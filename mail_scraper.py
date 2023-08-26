import os
import fitz  # PyMuPDF
import re
from docx import Document

def extract_emails_from_pdf(pdf_path):
    emails = set()  # Using a set to ensure uniqueness
    
    try:
        pdf_document = fitz.open(pdf_path)
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            text = page.get_text("text")
            email_matches = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b', text)
            emails.update(email_matches)  # Add email matches to the set
    except Exception as e:
        print("Error:", e)
    finally:
        pdf_document.close()
    
    return emails

def save_emails_to_word(emails, output_path):
    document = Document()
    document.add_heading("Extracted Emails", level=1)
    
    for email in emails:
        document.add_paragraph(email)
    
    document.save(output_path)

if __name__ == "__main__":
    pdf_directory = input("Enter the path of the directory containing pdf : ")
    output_path = "output_emails.docx"
    all_extracted_emails = set()  # Using a set to ensure uniqueness
    
    for pdf_filename in os.listdir(pdf_directory):
        if pdf_filename.endswith(".pdf"):
            pdf_path = os.path.join(pdf_directory, pdf_filename)
            extracted_emails = extract_emails_from_pdf(pdf_path)
            all_extracted_emails.update(extracted_emails)  # Add extracted emails to the set
    
    if all_extracted_emails:
        save_emails_to_word(all_extracted_emails, output_path)
        print(f"Emails saved to {output_path}")
    else:
        print("No emails found in the PDFs.")
